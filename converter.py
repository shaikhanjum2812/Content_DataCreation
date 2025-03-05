import logging
from docx import Document
from openpyxl import Workbook
import pandas as pd
import os
import tempfile
import shutil
from zipfile import ZipFile

# Set up logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def extract_sheet2_data(doc):
    """Extract QA data from Word document"""
    questions_data = []
    current_exid = None
    question_key = 1
    in_qa_section = False

    logger.info("Starting QA data extraction...")

    # First find EXID
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("exid :"):
            current_exid = text.split("exid :")[1].strip()
            logger.info(f"Found EXID: {current_exid}")
            break

    if not current_exid:
        logger.error("No EXID found in document")
        return []

    # Now extract QA pairs
    i = 0
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()

        if not text:
            i += 1
            continue

        if "Answer the following questions:" in text:
            in_qa_section = True
            logger.debug("Entered QA section")
            i += 1
            continue

        if in_qa_section and text.lower().startswith("question"):
            try:
                # Get question text
                question_parts = text.split(":", 1)
                question_text = question_parts[1].strip() if len(question_parts) > 1 else text
                logger.debug(f"Processing question: {question_text}")

                # Look ahead for answer
                if i + 1 < len(doc.paragraphs):
                    answer_text = doc.paragraphs[i + 1].text.strip()
                    logger.debug(f"Processing answer text: {answer_text}")

                    if "Options:" in answer_text and "answer:" in answer_text:
                        # Handle MCQ format
                        options_part = answer_text.split("Options:")[1]
                        options, answer_part = options_part.split("answer:")

                        options = options.strip()
                        answer = answer_part.strip()
                        hint = ""

                        if "Hint:" in answer:
                            answer, hint = answer.split("Hint:")
                            answer = answer.strip()
                            hint = hint.strip()

                        answer_type = "checkbox" if "," in answer else "radio"
                        if answer_type == "radio":
                            try:
                                answer = int(answer)
                            except ValueError:
                                pass

                        qa_entry = [
                            current_exid,
                            question_key,
                            question_text,
                            answer_type,
                            options,
                            answer,
                            hint
                        ]
                        questions_data.append(qa_entry)
                        logger.info(f"Added MCQ - Key: {question_key}")

                    elif "Answer:" in answer_text:
                        # Handle direct answer format
                        answer_part = answer_text.split("Answer:")[1]
                        answer = answer_part.strip()
                        hint = ""

                        if "Hint:" in answer:
                            answer, hint = answer.split("Hint:")
                            answer = answer.strip()
                            hint = hint.strip()

                        try:
                            answer = float(answer)
                            answer_type = "number"
                            if answer.is_integer():
                                answer = int(answer)
                        except ValueError:
                            answer_type = "text"

                        qa_entry = [
                            current_exid,
                            question_key,
                            question_text,
                            answer_type,
                            "",
                            answer,
                            hint
                        ]
                        questions_data.append(qa_entry)
                        logger.info(f"Added Direct Answer - Key: {question_key}")

                    question_key += 1
                    i += 2  # Skip the answer line
                    continue

            except Exception as e:
                logger.error(f"Error processing QA pair: {str(e)}")
                logger.error(f"Question text: {text}")
                if i + 1 < len(doc.paragraphs):
                    logger.error(f"Answer text: {doc.paragraphs[i + 1].text.strip()}")

        i += 1

    logger.info(f"Completed QA extraction. Total questions: {len(questions_data)}")
    for qa in questions_data:
        logger.debug(f"Extracted QA: {qa}")

    return questions_data

def extract_sheet1_data(doc):
    """Extract exercise metadata from Word document"""
    exercises_data = []
    current_data = {
        'exid': '', 'title': '', 'description': '', 'category': '',
        'subcategoryid': '', 'level': 0, 'language': '', 'qlocation': '',
        'module': '', 'ex_seq': 0, 'cat_seq': 0, 'subcat_seq': 0,
        'league': '', 'labels': ''
    }

    logger.info("Starting exercise data extraction...")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        for field in current_data.keys():
            prefix = f"{field} :"
            if text.startswith(prefix):
                value = text.split(prefix)[1].strip()
                if field in ['level', 'ex_seq', 'cat_seq', 'subcat_seq']:
                    try:
                        value = int(value)
                    except ValueError:
                        value = 0
                current_data[field] = value
                logger.debug(f"Extracted {field}: {value}")

                if field == 'labels':
                    exercises_data.append(list(current_data.values()))
                    logger.info(f"Added exercise data: {current_data['exid']}")
                    current_data = {k: '' if isinstance(v, str) else 0
                                    for k, v in current_data.items()}
                break

    logger.info(f"Total exercises extracted: {len(exercises_data)}")
    return exercises_data

def create_test_document():
    """Creates a test document with proper structure for data extraction"""
    doc = Document()
    logger.info("Creating test document with exercise data and QA sections")

    # Add exercise metadata with correct formatting
    doc.add_paragraph("exid : TEST001")
    doc.add_paragraph("title : Test Exercise")
    doc.add_paragraph("description : This is a test exercise for data extraction")
    doc.add_paragraph("category : Testing")
    doc.add_paragraph("subcategoryid : TEST")
    doc.add_paragraph("level : 1")
    doc.add_paragraph("language : python")
    doc.add_paragraph("qlocation : test.txt")
    doc.add_paragraph("module : test")
    doc.add_paragraph("ex_seq : 1")
    doc.add_paragraph("cat_seq : 1")
    doc.add_paragraph("subcat_seq : 1")
    doc.add_paragraph("league : beginner")
    doc.add_paragraph("labels : test,example")

    # Add code section
    doc.add_paragraph("Code:")
    doc.add_paragraph("def hello_world():")
    doc.add_paragraph("    print('Hello, World!')")
    doc.add_paragraph("    return True")

    # Add QA section marker
    doc.add_paragraph("Answer the following questions:")

    # Add questions with different formats
    # MCQ with hint
    doc.add_paragraph("Question 1: What will be the output of hello_world()?")
    doc.add_paragraph("Options: Hello World,Hi World,Hello, World!,World Hello answer: 3 Hint: Look at the print statement")

    # Checkbox with hint
    doc.add_paragraph("Question 2: What Python concepts are used in the code?")
    doc.add_paragraph("Options: function definition,print statement,return statement,variables answer: 1,2,3 Hint: Look at code structure")

    # Number answer with hint
    doc.add_paragraph("Question 3: How many lines of code are in the function?")
    doc.add_paragraph("Answer: 2 Hint: Count the indented lines")

    # Text answer with hint
    doc.add_paragraph("Question 4: What is the purpose of the return statement?")
    doc.add_paragraph("Answer: To indicate successful execution Hint: Think about function behavior")

    # Save test document
    test_path = "test_document.docx"
    doc.save(test_path)
    logger.info(f"Created test document at {test_path}")
    return test_path

def convert_word_to_excel(input_path, output_path):
    """Convert Word document to Excel format with exercise and QA data"""
    try:
        logger.info(f"Starting conversion of {input_path}")
        doc = Document(input_path)

        # Extract data
        sheet1_data = extract_sheet1_data(doc)
        sheet2_data = extract_sheet2_data(doc)

        # Verify extracted data
        logger.info("\nExtraction Summary:")
        logger.info(f"Number of exercises: {len(sheet1_data)}")
        logger.info(f"Number of questions: {len(sheet2_data)}")

        if not sheet1_data:
            logger.error("No exercise data extracted!")
            return False

        if not sheet2_data:
            logger.error("No QA data extracted!")
            return False

        # Create DataFrames with explicit data
        columns_sheet1 = ['exid', 'title', 'description', 'category', 
                         'subcategoryid', 'level', 'language', 'qlocation', 
                         'module', 'ex_seq', 'cat_seq', 'subcat_seq', 
                         'league', 'labels']
        columns_sheet2 = ['exid', 'key', 'question', 'type', 'options', 
                         'answer', 'hint']

        # Convert data to pandas DataFrame with explicit columns
        df_sheet1 = pd.DataFrame(sheet1_data, columns=columns_sheet1)
        df_sheet2 = pd.DataFrame(sheet2_data, columns=columns_sheet2)

        # Validate DataFrames
        if df_sheet1.empty or df_sheet2.empty:
            logger.error("One or both DataFrames are empty!")
            return False

        # Write to Excel file with proper formatting
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            # Write Sheet 1
            df_sheet1.to_excel(writer, sheet_name='ex_data', index=False)
            worksheet1 = writer.sheets['ex_data']
            for column in worksheet1.columns:
                max_length = 0
                column = [cell for cell in column]
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(cell.value)
                    except:
                        pass
                adjusted_width = (max_length + 2)
                worksheet1.column_dimensions[column[0].column_letter].width = adjusted_width

            # Write Sheet 2
            df_sheet2.to_excel(writer, sheet_name='qa_data', index=False)
            worksheet2 = writer.sheets['qa_data']
            for column in worksheet2.columns:
                max_length = 0
                column = [cell for cell in column]
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(cell.value)
                    except:
                        pass
                adjusted_width = (max_length + 2)
                worksheet2.column_dimensions[column[0].column_letter].width = adjusted_width

        logger.info(f"Successfully wrote data to Excel: {output_path}")
        logger.info(f"Sheet1 rows: {len(df_sheet1)}")
        logger.info(f"Sheet2 rows: {len(df_sheet2)}")

        # Verify file was created and has content
        if not os.path.exists(output_path):
            logger.error("Excel file was not created!")
            return False

        if os.path.getsize(output_path) == 0:
            logger.error("Excel file is empty!")
            return False

        return True

    except Exception as e:
        logger.error(f"Error converting file: {str(e)}")
        logger.error("Stack trace:", exc_info=True)
        raise

def create_text_files(input_path):
    """
    Creates text files from code blocks in a Word document and returns a zip file path.
    """
    try:
        doc = Document(input_path)
        temp_dir = tempfile.mkdtemp()

        # Extract code blocks
        queries = []
        collecting_code = False
        current_code = []
        qlocation = None

        for para in doc.paragraphs:
            text = para.text.strip()

            if text.startswith("qlocation :"):
                qlocation = text.split("qlocation :")[1].strip()
                if not qlocation.endswith('.txt'):
                    qlocation = f"{qlocation}.txt"

            elif text.startswith("Code:"):
                collecting_code = True
                current_code = []

            elif collecting_code and "Answer the following questions:" in text:
                if current_code and qlocation:
                    queries.append({
                        "qlocation": qlocation,
                        "code": "\n".join(current_code)
                    })
                collecting_code = False
                current_code = []

            elif collecting_code:
                current_code.append(text)

        # Create text files and add to zip
        zip_path = os.path.join(tempfile.gettempdir(), 'code_files.zip')
        with ZipFile(zip_path, 'w') as zipf:
            for query in queries:
                file_path = os.path.join(temp_dir, query["qlocation"])
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(query["code"])
                zipf.write(file_path, query["qlocation"])

        shutil.rmtree(temp_dir)
        return zip_path

    except Exception as e:
        logger.error(f"Error creating text files: {str(e)}")
        raise


if __name__ == "__main__":
    # Create output directory if it doesn't exist
    if not os.path.exists("output"):
        os.makedirs("output")

    try:
        # Create and process test document
        test_doc_path = create_test_document()
        output_excel = "output/test_output.xlsx"
        success = convert_word_to_excel(test_doc_path, output_excel)

        if success:
            logger.info(f"Successfully converted document to {output_excel}")
            os.remove(test_doc_path)
            logger.info("Test completed successfully")
        else:
            logger.error("Conversion failed")

    except Exception as e:
        logger.error(f"Test failed with error: {str(e)}")