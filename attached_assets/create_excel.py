import os
import pandas as pd
from docx import Document

# Function to extract data from the Word document for Sheet1
def extract_sheet1_data_from_docx(docx_path):
    document = Document(docx_path)
    data = []
    
    current_exid = ""
    current_title = ""
    current_description = ""
    current_category = ""
    current_subcategoryid = ""
    current_level = 0
    current_language = ""
    current_qlocation = ""
    current_module = ""
    current_ex_seq = 0
    current_cat_seq = 0
    current_subcat_seq = 0
    current_league = ""
    current_labels = ""
    
    for para in document.paragraphs:
        text = para.text.strip()
        
        if text.startswith("exid :"):
            if current_exid:
                data.append([
                    current_exid, current_title, current_description, current_category,
                    current_subcategoryid, current_level, current_language, current_qlocation,
                    current_module, current_ex_seq, current_cat_seq, current_subcat_seq,
                    current_league, current_labels
                ])
            current_exid = text.split("exid :")[1].strip()
        elif text.startswith("title :"):
            current_title = text.split("title :")[1].strip()
        elif text.startswith("description :"):
            current_description = text.split("description :")[1].strip()
        elif text.startswith("category :"):
            current_category = text.split("category :")[1].strip()
        elif text.startswith("subcategoryid :"):
            current_subcategoryid = text.split("subcategoryid :")[1].strip()
        elif text.startswith("level :"):
            current_level = int(text.split("level :")[1].strip())
        elif text.startswith("language :"):
            current_language = text.split("language :")[1].strip()
        elif text.startswith("qlocation :"):
            current_qlocation = text.split("qlocation :")[1].strip()
        elif text.startswith("module :"):
            current_module = text.split("module :")[1].strip()
        elif text.startswith("ex_seq :"):
            current_ex_seq = int(text.split("ex_seq :")[1].strip())
        elif text.startswith("cat_seq :"):
            current_cat_seq = int(text.split("cat_seq :")[1].strip())
        elif text.startswith("subcat_seq :"):
            current_subcat_seq = int(text.split("subcat_seq :")[1].strip())
        elif text.startswith("league :"):
            current_league = text.split("league :")[1].strip()
        elif text.startswith("labels :"):
            current_labels = text.split("labels :")[1].strip()
    
    if current_exid:
        data.append([
            current_exid, current_title, current_description, current_category,
            current_subcategoryid, current_level, current_language, current_qlocation,
            current_module, current_ex_seq, current_cat_seq, current_subcat_seq,
            current_league, current_labels
        ])
    
    return data

# Function to extract questions and answers from the Word document for Sheet2
def extract_sheet2_data_from_docx(docx_path):
    document = Document(docx_path)
    questions_data = []
    exid = ""
    question_key = 1
    
    for para in document.paragraphs:
        text = para.text.strip()
        
        if text.startswith("exid :"):
            exid = text.split("exid :")[1].strip()
            question_key = 1
        elif text.startswith("Answer the following questions:"):
            continue
        elif "Options:" in text:
            question, options_answer = text.split("Options:")
            options, answer = options_answer.split("Answer:")
            options = options.strip().split(',')
            answer = answer.strip()
            if ',' in answer:
                answer_type = "checkbox"
            else:
                answer_type = "radio"
                answer = int(answer)
            questions_data.append([exid, question_key, question.strip(), answer_type, ','.join(options), answer])
            question_key += 1
        elif "Answer:" in text:
            question, answer = text.split("Answer:")
            answer = answer.strip()
            try:
                answer = float(answer)
                answer_type = "number"
                if answer.is_integer():
                    answer = int(answer)
            except ValueError:
                answer_type = "text"
            questions_data.append([exid, question_key, question.strip(), answer_type, "", answer])
            question_key += 1
            
    return questions_data

# Paths to the Word document and the Excel file
docx_path = "Z:/DATASCIENCELAB/Projects/Pravinyam_Project/CProgramFiles/For each loop.docx"
excel_path = "Z:/DATASCIENCELAB/Projects/Pravinyam_Project/CProgramFiles/For each loop1.xlsx"
new_excel_path = "Z:/DATASCIENCELAB/Projects/Pravinyam_Project/CProgramFiles/new_excel.xlsx"

# Extract data from the Word document
sheet1_data = extract_sheet1_data_from_docx(docx_path)
sheet2_data = extract_sheet2_data_from_docx(docx_path)

# Create DataFrames for Sheet1 and Sheet2
columns_sheet1 = ['exid', 'title', 'description', 'category', 'subcategoryid', 'level', 'language', 'qlocation', 'module', 'ex_seq', 'cat_seq', 'subcat_seq', 'league', 'labels']
columns_sheet2 = ['exid', 'key', 'label', 'type', 'options', 'answer']

df_sheet1 = pd.DataFrame(sheet1_data, columns=columns_sheet1)
df_sheet2 = pd.DataFrame(sheet2_data, columns=columns_sheet2)

# Check if the Excel file exists, if not create a new one
if not os.path.exists(excel_path):
    with pd.ExcelWriter(new_excel_path, engine='openpyxl') as writer:
        df_sheet1.to_excel(writer, sheet_name='ex_data', index=False)
        df_sheet2.to_excel(writer, sheet_name='qa_data', index=False)
    print(f"New Excel file created at {new_excel_path}")
else:
    with pd.ExcelWriter(excel_path, engine='openpyxl', mode='a', if_sheet_exists='replace') as writer:
        df_sheet1.to_excel(writer, sheet_name='ex_data', index=False)
        df_sheet2.to_excel(writer, sheet_name='qa_data', index=False)
    print("Excel file updated successfully!")
