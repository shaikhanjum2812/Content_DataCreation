import os
import logging
from docx import Document
from converter import convert_word_to_excel

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def create_test_document():
    """Creates a test document with sample data"""
    doc = Document()

    # Add exercise metadata
    doc.add_paragraph("exid : TEST001")
    doc.add_paragraph("title : Test Exercise")
    doc.add_paragraph("description : This is a test exercise")
    doc.add_paragraph("category : Testing")
    doc.add_paragraph("subcategoryid : TEST")
    doc.add_paragraph("level : 1")
    doc.add_paragraph("language : python")
    doc.add_paragraph("qlocation : test.txt")
    doc.add_paragraph("module : test")
    doc.add_paragraph("ex_seq : 1")
    doc.add_paragraph("cat_seq : 1")
    doc.add_paragraph("subcat_seq : 1")
    doc.add_paragraph("league : beginner")
    doc.add_paragraph("labels : test,example")

    # Add questions section
    doc.add_paragraph("Answer the following questions:")

    # Add multiple choice question
    doc.add_paragraph("What is Python?")
    doc.add_paragraph("Options: A programming language,An animal,A snake,A food item answer: 1")

    # Add checkbox question
    doc.add_paragraph("Select all that apply to Python:")
    doc.add_paragraph("Options: Interpreted,Compiled,Object-oriented,Functional answer: 1,3,4")

    # Add a direct answer question
    doc.add_paragraph("How many bits are in a byte?")
    doc.add_paragraph("Answer: 8")

    # Save the test document
    test_path = "test_document.docx"
    doc.save(test_path)
    logger.info(f"Created test document: {test_path}")
    return test_path

def verify_excel_output(excel_path):
    """Verify that the Excel file exists and has content"""
    if not os.path.exists(excel_path):
        logger.error(f"Excel file not found: {excel_path}")
        return False

    if os.path.getsize(excel_path) == 0:
        logger.error("Excel file is empty")
        return False

    logger.info(f"Excel file verified: {excel_path}")
    return True

def main():
    try:
        # Create output directory if it doesn't exist
        if not os.path.exists("output"):
            os.makedirs("output")

        # Create and process test document
        test_doc = create_test_document()
        logger.info(f"Test document created: {test_doc}")

        # Convert to Excel
        output_excel = "output/test_output.xlsx"
        success = convert_word_to_excel(test_doc, output_excel)

        if success and verify_excel_output(output_excel):
            logger.info("Test completed successfully")
            logger.info(f"Output file: {output_excel}")
        else:
            logger.error("Test failed - conversion or verification failed")

    except Exception as e:
        logger.error(f"Test failed with error: {str(e)}")
    finally:
        # Clean up test document
        if os.path.exists(test_doc):
            os.remove(test_doc)
            logger.info("Test document cleaned up")

if __name__ == "__main__":
    main()